"""
DOCX Builder — Professional RICS Report Word Document Generator
Converts the HTML report to a styled Word document with embedded images.

Uses python-docx to create a RICS-branded DOCX file that surveyors
can freely edit, annotate, and redistribute.
"""

import os
import re
import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger("docx_builder")

# RICS brand colors (in RGBColor format for python-docx)
PURPLE_HEX = "4D2D69"
GOLD_HEX = "D4A843"
GREEN_HEX = "4CAF50"
AMBER_HEX = "FF9800"
RED_HEX = "F44336"
DARK_HEX = "231F20"
GRAY_HEX = "6B7280"


def _hex_to_rgb(hex_str: str):
    """Convert hex color string to docx RGBColor."""
    from docx.shared import RGBColor
    return RGBColor(int(hex_str[:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def _add_image_safe(doc_or_cell, img_path: str, width_inches: float = 5.0):
    """Add image to document with error handling and resizing."""
    from docx.shared import Inches
    try:
        if not os.path.exists(img_path):
            logger.warning(f"Image not found: {img_path}")
            return None

        # Resize with Pillow if available
        try:
            from PIL import Image
            img = Image.open(img_path)
            max_w = int(width_inches * 150)  # ~150 DPI target
            if img.width > max_w:
                ratio = max_w / img.width
                new_size = (max_w, int(img.height * ratio))
                img = img.resize(new_size, Image.LANCZOS)

            buf = io.BytesIO()
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            img.save(buf, format='JPEG', quality=80, optimize=True)
            buf.seek(0)
            return doc_or_cell.add_picture(buf, width=Inches(width_inches))
        except ImportError:
            # Fall back to adding raw file
            return doc_or_cell.add_picture(img_path, width=Inches(width_inches))
    except Exception as e:
        logger.error(f"Failed to add image {img_path}: {e}")
        return None


def _resolve_img_path(src: str, storage_root: str) -> str:
    """Resolve an image src (from HTML) to an absolute filesystem path."""
    if src.startswith("/storage/"):
        return os.path.join(storage_root, src.replace("/storage/", ""))
    if src.startswith("file://"):
        return src.replace("file://", "")
    if os.path.isabs(src):
        return src
    return os.path.join(storage_root, src)


def build_docx_from_html(
    html_path: str,
    output_path: str,
    reference: str = "",
    storage_root: str = "/app/storage",
) -> str:
    """
    Build a professional RICS Word document from the HTML report.

    Parses the HTML structure, extracts text and images,
    and creates a properly styled DOCX file.

    Args:
        html_path: Path to the HTML report file
        output_path: Where to save the DOCX
        reference: Project reference number
        storage_root: Root directory for resolving image paths

    Returns:
        Path to generated DOCX
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    logger.info(f"Building DOCX from: {html_path}")

    # Read HTML
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = _hex_to_rgb(DARK_HEX)
    style.paragraph_format.space_after = Pt(6)

    # ── Parse HTML for structure ──
    # We'll use a simplified parser that handles the RICS report structure
    _parse_and_build(doc, html, storage_root)

    # ── Footer ──
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"{reference}  |  RICS Home Survey – Level 3")
    run.font.size = Pt(8)
    run.font.color.rgb = _hex_to_rgb(GRAY_HEX)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    size_kb = os.path.getsize(output_path) // 1024
    logger.info(f"DOCX generated: {output_path} ({size_kb}KB)")
    return output_path


def _parse_and_build(doc, html: str, storage_root: str):
    """
    Parse HTML content and build DOCX structure.
    Handles headings, paragraphs, images, tables, and the RICS rating system.
    """
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup

    # Try BeautifulSoup first, fallback to regex
    try:
        soup = BeautifulSoup(html, 'html.parser')
        _build_from_soup(doc, soup, storage_root)
    except ImportError:
        logger.info("BeautifulSoup not available, using regex parser")
        _build_from_regex(doc, html, storage_root)


def _build_from_soup(doc, soup, storage_root: str):
    """Build DOCX from BeautifulSoup parsed HTML."""
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    body = soup.find('body') or soup

    # Find the main content area
    content = body.find('div', class_='report-body') or body

    for element in content.children:
        if not hasattr(element, 'name') or element.name is None:
            # Text node
            text = element.strip() if isinstance(element, str) else ''
            if text:
                doc.add_paragraph(text)
            continue

        tag = element.name.lower()

        # ── Cover page elements ──
        if 'cover' in (element.get('class', []) or []) or 'report-cover' in str(element.get('class', '')):
            _add_cover_section(doc, element)
            doc.add_page_break()
            continue

        # ── Headings ──
        if tag in ('h1', 'h2', 'h3', 'h4'):
            level = int(tag[1])
            text = element.get_text(strip=True)
            if text:
                p = doc.add_heading(text, level=level)
                # Purple color for main headings
                if level <= 2:
                    for run in p.runs:
                        run.font.color.rgb = _hex_to_rgb(PURPLE_HEX)
            continue

        # ── Tables ──
        if tag == 'table':
            _add_table_from_soup(doc, element)
            continue

        # ── Images ──
        if tag == 'img':
            src = element.get('src', '')
            if src:
                path = _resolve_img_path(src, storage_root)
                _add_image_safe(doc, path, width_inches=5.5)
            continue

        # ── Divs with special classes ──
        classes = element.get('class', []) or []
        class_str = ' '.join(classes) if isinstance(classes, list) else str(classes)

        if 'evidence-grid' in class_str or 'photo-grid' in class_str:
            _add_evidence_section(doc, element, storage_root)
            continue

        if 'condition-badge' in class_str or 'rating' in class_str:
            _add_rating_badge(doc, element)
            continue

        if 'section-header' in class_str:
            text = element.get_text(strip=True)
            if text:
                p = doc.add_heading(text, level=2)
                for run in p.runs:
                    run.font.color.rgb = _hex_to_rgb(PURPLE_HEX)
            continue

        # ── Generic div/section — recurse ──
        if tag in ('div', 'section', 'article', 'main', 'header', 'footer', 'aside'):
            _build_from_soup(doc, element, storage_root)
            continue

        # ── Paragraphs ──
        if tag == 'p':
            text = element.get_text(strip=True)
            if text:
                p = doc.add_paragraph(text)
                # Check for bold
                if element.find('strong') or element.find('b'):
                    for run in p.runs:
                        run.bold = True
            continue

        # ── Lists ──
        if tag in ('ul', 'ol'):
            for li in element.find_all('li', recursive=False):
                text = li.get_text(strip=True)
                if text:
                    doc.add_paragraph(text, style='List Bullet')
            continue

        # ── Horizontal rule ──
        if tag == 'hr':
            p = doc.add_paragraph()
            p.add_run('─' * 60).font.color.rgb = _hex_to_rgb(GRAY_HEX)
            continue

        # Fallback: try to get text
        text = element.get_text(strip=True)
        if text and len(text) > 3:
            doc.add_paragraph(text)


def _add_cover_section(doc, element):
    """Add cover page content."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = element.get_text('\n', strip=True)
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        if i == 0:
            run.font.size = Pt(28)
            run.font.color.rgb = _hex_to_rgb(PURPLE_HEX)
            run.bold = True
        elif i < 3:
            run.font.size = Pt(16)
            run.font.color.rgb = _hex_to_rgb(GOLD_HEX)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = _hex_to_rgb(GRAY_HEX)


def _add_table_from_soup(doc, table_el):
    """Convert HTML table to DOCX table."""
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT

    rows = table_el.find_all('tr')
    if not rows:
        return

    max_cols = max(len(r.find_all(['td', 'th'])) for r in rows) if rows else 0
    if max_cols == 0:
        return

    table = doc.add_table(rows=0, cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_el in rows:
        cells = row_el.find_all(['td', 'th'])
        row = table.add_row()
        for i, cell_el in enumerate(cells):
            if i < max_cols:
                text = cell_el.get_text(strip=True)
                row.cells[i].text = text
                # Header styling
                if cell_el.name == 'th':
                    for paragraph in row.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)


def _add_evidence_section(doc, element, storage_root: str):
    """Add photo evidence grid to DOCX."""
    from docx.shared import Inches

    imgs = element.find_all('img')
    for img in imgs:
        src = img.get('src', '')
        if src:
            path = _resolve_img_path(src, storage_root)
            _add_image_safe(doc, path, width_inches=4.5)

            # Check for caption/annotation
            caption = img.get('alt', '') or img.get('title', '')
            parent = img.parent
            if parent:
                ann = parent.find(['p', 'span', 'div'], class_=lambda c: c and ('caption' in str(c) or 'annotation' in str(c) or 'defect' in str(c)))
                if ann:
                    caption = ann.get_text(strip=True)

            if caption:
                p = doc.add_paragraph(caption)
                p.style = 'Caption' if 'Caption' in [s.name for s in doc.styles] else 'Normal'


def _add_rating_badge(doc, element):
    """Add condition rating badge as styled text."""
    from docx.shared import Pt
    text = element.get_text(strip=True)
    if not text:
        return

    p = doc.add_paragraph()
    run = p.add_run(f"[{text}]")
    run.bold = True
    run.font.size = Pt(11)

    t = text.lower()
    if '3' in t or 'urgent' in t or 'red' in t:
        run.font.color.rgb = _hex_to_rgb(RED_HEX)
    elif '2' in t or 'amber' in t or 'attention' in t:
        run.font.color.rgb = _hex_to_rgb(AMBER_HEX)
    else:
        run.font.color.rgb = _hex_to_rgb(GREEN_HEX)


def _build_from_regex(doc, html: str, storage_root: str):
    """
    Fallback parser using regex when BeautifulSoup is not available.
    Handles the most common HTML patterns in RICS reports.
    """
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Strip HTML tags but preserve structure
    # Remove script/style
    html = re.sub(r'<(script|style)[^>]*>.*?</\1>', '', html, flags=re.DOTALL | re.IGNORECASE)

    # Headings
    for m in re.finditer(r'<h([1-4])[^>]*>(.*?)</h\1>', html, re.DOTALL | re.IGNORECASE):
        level = int(m.group(1))
        text = re.sub(r'<[^>]+>', '', m.group(2)).strip()
        if text:
            p = doc.add_heading(text, level=level)
            if level <= 2:
                for run in p.runs:
                    run.font.color.rgb = _hex_to_rgb(PURPLE_HEX)

    # Images
    for m in re.finditer(r'<img[^>]+src=["\']([^"\']+)["\']', html, re.IGNORECASE):
        src = m.group(1)
        path = _resolve_img_path(src, storage_root)
        _add_image_safe(doc, path, width_inches=5.0)

    # Paragraphs
    for m in re.finditer(r'<p[^>]*>(.*?)</p>', html, re.DOTALL | re.IGNORECASE):
        text = re.sub(r'<[^>]+>', '', m.group(1)).strip()
        if text and len(text) > 2:
            doc.add_paragraph(text)


def build_room_docx(
    partial_report_path: str,
    room_folder: str,
    output_path: str,
    reference: str = "",
    storage_root: str = "/app/storage",
) -> str:
    """
    Build a DOCX for a single room from its partial_report.json and photos.

    Args:
        partial_report_path: Path to partial_report.json
        room_folder: Path to the room's folder (containing Context_* dirs)
        output_path: Where to save the DOCX
        reference: Project reference
        storage_root: Storage root for resolving paths

    Returns:
        Path to generated DOCX
    """
    import json
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with open(partial_report_path, "r", encoding="utf-8") as f:
        report = json.load(f)

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Title
    p = doc.add_heading(f"Room Report: {report.get('room_name', 'Unknown')}", level=1)
    for run in p.runs:
        run.font.color.rgb = _hex_to_rgb(PURPLE_HEX)

    # Meta
    doc.add_paragraph(f"Reference: {reference}")
    doc.add_paragraph(f"Floor: {report.get('floor_level', 'N/A')}")
    doc.add_paragraph("")

    # Summary
    summary = report.get('inspection_summary', '')
    if summary:
        doc.add_heading("Inspection Summary", level=2)
        doc.add_paragraph(summary)

    # Elements
    elements = report.get('elements', [])
    for el in elements:
        doc.add_heading(el.get('rics_element', el.get('name', '')), level=3)

        # Condition rating
        cr = el.get('condition_rating', 0)
        cr_text = {1: 'CR1 — No repair needed', 2: 'CR2 — Repair needed', 3: 'CR3 — Urgent repair'}
        p = doc.add_paragraph()
        run = p.add_run(f"Condition Rating: {cr_text.get(cr, f'CR{cr}')}")
        run.bold = True
        color = {1: GREEN_HEX, 2: AMBER_HEX, 3: RED_HEX}.get(cr, GRAY_HEX)
        run.font.color.rgb = _hex_to_rgb(color)

        # Description
        desc = el.get('condition_description', '')
        if desc:
            doc.add_paragraph(desc)

        # Defects
        defects = el.get('defects_identified', [])
        for d in defects:
            p = doc.add_paragraph()
            run = p.add_run(f"⚠ {d.get('defect_type', 'Defect')}")
            run.bold = True
            run.font.color.rgb = _hex_to_rgb(AMBER_HEX)
            if d.get('severity'):
                doc.add_paragraph(f"  Severity: {d['severity']}")
            if d.get('location'):
                doc.add_paragraph(f"  Location: {d['location']}")
            if d.get('probable_cause'):
                doc.add_paragraph(f"  Cause: {d['probable_cause']}")
            if d.get('recommended_action'):
                doc.add_paragraph(f"  Action: {d['recommended_action']}")

        # Evidence photos
        photos = el.get('evidence_photos', [])
        if photos:
            doc.add_heading("Evidence Photos", level=4)
            for photo_url in photos:
                path = _resolve_img_path(photo_url, storage_root)
                _add_image_safe(doc, path, width_inches=4.5)

    # Add remaining photos from context folders
    if os.path.isdir(room_folder):
        doc.add_page_break()
        doc.add_heading("All Context Photos", level=2)
        for ctx_name in sorted(os.listdir(room_folder)):
            ctx_path = os.path.join(room_folder, ctx_name)
            if not os.path.isdir(ctx_path) or not ctx_name.startswith("Context_"):
                continue

            label = ctx_name.replace("Context_", "").replace("_", " ")
            doc.add_heading(f"📷 {label}", level=3)

            for fname in sorted(os.listdir(ctx_path)):
                if fname.lower().endswith(('.jpg', '.jpeg', '.png')):
                    img_path = os.path.join(ctx_path, fname)
                    _add_image_safe(doc, img_path, width_inches=4.5)

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"{reference} | Room: {report.get('room_name', '')} | RICS Level 3")
    run.font.size = Pt(8)
    run.font.color.rgb = _hex_to_rgb(GRAY_HEX)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    logger.info(f"Room DOCX generated: {output_path}")
    return output_path
